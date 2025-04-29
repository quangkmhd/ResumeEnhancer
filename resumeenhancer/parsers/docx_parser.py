"""
Module đọc và phân tích file DOCX.
"""
from pathlib import Path
from typing import Optional
import docx
from loguru import logger


def extract_text_from_docx(docx_path: str) -> Optional[str]:
    """
    Trích xuất nội dung văn bản từ file DOCX.
    
    Args:
        docx_path: Đường dẫn tới file DOCX.
    
    Returns:
        Nội dung văn bản của file DOCX hoặc None nếu xảy ra lỗi.
    """
    try:
        logger.info(f"Đang đọc file DOCX: {docx_path}")
        
        # Kiểm tra file có tồn tại không
        docx_file = Path(docx_path)
        if not docx_file.exists():
            logger.error(f"File không tồn tại: {docx_path}")
            return None
        
        # Đọc nội dung DOCX
        doc = docx.Document(docx_path)
        text = ""
        
        # Trích xuất văn bản từ tất cả các đoạn
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
        
        # Trích xuất văn bản từ các bảng
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text += cell.text + " "
                text += "\n"
        
        if not text.strip():
            logger.warning(f"Không trích xuất được văn bản từ file DOCX: {docx_path}")
            return None
        
        logger.debug(f"Đã trích xuất thành công {len(text)} ký tự từ file DOCX")
        return text
        
    except Exception as e:
        logger.error(f"Lỗi khi đọc file DOCX {docx_path}: {str(e)}")
        return None 